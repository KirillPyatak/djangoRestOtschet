# report_app/report_generators.py
from docx import Document
from openpyxl import Workbook

class WordReportGenerator:
    def generate(self, queryset):
        document = Document()
        for item in queryset:
            document.add_paragraph(f'Field1: {item.field1}')
            document.add_paragraph(f'Field2: {item.field2}')
            # Добавьте другие поля, если они есть
        return document

class ExcelReportGenerator:
    def generate(self, queryset):
        wb = Workbook()
        ws = wb.active
        for item in queryset:
            ws.append([item.field1, item.field2])
            # Добавьте другие поля, если они есть
        return wb
