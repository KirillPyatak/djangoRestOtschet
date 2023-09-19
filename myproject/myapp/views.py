from django.shortcuts import render
from django.http import HttpResponse
from docx import Document
from openpyxl import Workbook
import tempfile
import os

from .models import ReportData  # Импортируйте модель данных из вашего приложения

# Класс для работы с отчетами в формате Word
class WordReportGenerator:
    def generate(self, queryset):
        document = Document()
        for item in queryset:
            document.add_paragraph(f'Field1: {item.field1}')
            document.add_paragraph(f'Field2: {item.field2}')
            # Добавьте другие поля, если они есть
        return document

# Класс для работы с отчетами в формате Excel
class ExcelReportGenerator:
    def generate(self, queryset):
        wb = Workbook()
        ws = wb.active
        for item in queryset:
            ws.append([item.field1, item.field2])
            # Добавьте другие поля, если они есть
        return wb

# Функция для генерации и скачивания отчета
def generate_report(request):
    if request.method == 'POST':
        format = request.POST.get('format')

        # Получите данные из базы данных
        queryset = ReportData.objects.all()

        try:
            if format == 'word':
                generator = WordReportGenerator()
                document = generator.generate(queryset)
                temp_file_path = save_temp_document(document, "docx")
                return download_temp_file(temp_file_path, "report.docx")
            elif format == 'excel':
                generator = ExcelReportGenerator()
                workbook = generator.generate(queryset)
                temp_file_path = save_temp_document(workbook, "xlsx")
                return download_temp_file(temp_file_path, "report.xlsx")
        except Exception as e:
            return HttpResponse("An error occurred while generating the report.", status=500)

    return render(request, 'index.html')

# Функция для сохранения временного файла
def save_temp_document(document, file_extension):
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_extension}")
    temp_file_path = temp_file.name
    document.save(temp_file_path)
    return temp_file_path

# Функция для скачивания временного файла
def download_temp_file(temp_file_path, filename):
    with open(temp_file_path, 'rb') as temp_file:
        response = HttpResponse(temp_file.read())
        response['Content-Disposition'] = f'attachment; filename={filename}'
        return response

# Добавьте логирование действий в приложении
def generate_report_page(request):
    return render(request, 'index.html')
